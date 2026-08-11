from __future__ import annotations

import base64
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
import pytest

from bookforge.contracts.assembly import AssemblyNotReadyError, BookMetadataV3
from bookforge.contracts.classification import (
    ClassificationProvenance, ClassificationResult, ClassifierIdentity,
    ClassifierKind, ReviewStatus,
)
from bookforge.contracts.flow import LogicalListKind, LogicalListV3, StructuralRegion, StructuralRegionAssignment
from bookforge.contracts.flow import ContinuityType, FlowDecisionReview, ResolverIdentity, ResolverKind
from bookforge.contracts.assembly import TextSemanticNode
from bookforge.flow.models import AcceptedFlowReviewInput
from bookforge.contracts.ids import classification_result_id
from bookforge.contracts.raw import RawImage
from bookforge.contracts.semantic import SemanticType
from bookforge.contracts.source import SourceTextReference
from bookforge.docx import DocxExtractor
from bookforge.flow.models import FlowSourceFeatures
from bookforge.pipeline import PipelineInput, PipelineIntegrationError, PipelineRunner
from bookforge.semantic.models import AnalysisView
from bookforge.semantic.pipeline import EPOCH, generate_work_units

FP = "a" * 64
PNG = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nT8AAAAASUVORK5CYII=")


class ExplicitFixtureClassifier:
    identity = ClassifierIdentity(
        name="m5c-explicit-fixture", kind=ClassifierKind.DETERMINISTIC, version="1"
    )

    def __init__(self, semantic_by_unit: dict[str, SemanticType]) -> None:
        self.semantic_by_unit = dict(semantic_by_unit)
        self.configuration_fingerprint = FP

    def classify(self, analysis_view: AnalysisView) -> ClassificationResult:
        unit = analysis_view.work_unit
        semantic_type = self.semantic_by_unit.get(unit.work_unit_id, SemanticType.UNKNOWN)
        references = ()
        if unit.source_kind.value in {"paragraph", "text_block"}:
            references = tuple(SourceTextReference(source_id=item) for item in unit.target_source_ids)
        result_id = classification_result_id(
            target_source_ids=[str(item) for item in unit.target_source_ids],
            taxonomy_version="bookforge-semantic-v1", classifier_name=self.identity.name,
            classifier_version=self.identity.version,
            configuration_fingerprint=self.configuration_fingerprint,
            input_fingerprint=unit.input_fingerprint,
            context_fingerprint=unit.context_fingerprint,
        )
        return ClassificationResult(
            id=result_id, target_source_ids=unit.target_source_ids,
            source_references=references, semantic_type=semantic_type,
            confidence=1 if semantic_type is not SemanticType.UNKNOWN else 0,
            review_status=ReviewStatus.NOT_REQUIRED if semantic_type is not SemanticType.UNKNOWN else ReviewStatus.NEEDS_REVIEW,
            classifier=self.identity, configuration_fingerprint=self.configuration_fingerprint,
            input_fingerprint=unit.input_fingerprint, context_fingerprint=unit.context_fingerprint,
            provenance=ClassificationProvenance(
                document_id=unit.document_id, source_ids=unit.target_source_ids, created_at=EPOCH
            ),
        )


def _make_docx(path: Path, image_path: Path) -> None:
    image_path.write_bytes(PNG)
    doc = Document()
    for text in (
        "Sách Việt & XML <đúng>", "Tác giả", "PHẦN I", "Chương Một",
        "Mục 1", "Đoạn tiếng Việt.",
    ):
        doc.add_paragraph(text)
    doc.add_picture(str(image_path), width=Inches(0.2))
    doc.add_paragraph("Hình 1 — minh họa")
    doc.add_paragraph("Trích dẫn A")
    doc.add_paragraph("Trích dẫn B")
    doc.add_paragraph("Mục lớn 1")
    doc.add_paragraph("Mục lớn 2")
    doc.add_paragraph("Mục con 2.1")
    doc.add_paragraph("Mục con 2.2")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "Ô A", "Ô B"
    table.cell(1, 0).text, table.cell(1, 1).text = "Ô C", "Ô D"
    doc.add_paragraph("Ghi chú cuối")
    doc.save(path)


def _fixture_truth(source: Path, root: Path):
    extraction = DocxExtractor().extract(source, root / "discovery")
    units = generate_work_units(extraction.raw_document)
    types = (
        SemanticType.BOOK_TITLE, SemanticType.AUTHOR, SemanticType.PART_TITLE,
        SemanticType.CHAPTER_HEADING, SemanticType.SECTION_HEADING,
        SemanticType.PARAGRAPH, SemanticType.QUOTE, SemanticType.FIGURE,
        SemanticType.CAPTION, SemanticType.QUOTE, SemanticType.QUOTE,
        SemanticType.LIST_ITEM, SemanticType.LIST_ITEM,
        SemanticType.LIST_ITEM, SemanticType.LIST_ITEM,
        SemanticType.TABLE, SemanticType.NOTE,
    )
    assert len(units) == len(types)
    classifier = ExplicitFixtureClassifier({unit.work_unit_id: kind for unit, kind in zip(units, types, strict=True)})
    fragment_ids = tuple(f"sem_f{index:06d}" for index in range(1, len(units) + 1))
    front = set(fragment_ids[:2])
    back = {fragment_ids[-1]}
    regions = StructuralRegionAssignment(by_fragment_id={
        item: StructuralRegion.FRONT if item in front else StructuralRegion.BACK if item in back else StructuralRegion.BODY
        for item in fragment_ids
    })
    parent = LogicalListV3(
        list_id="list_aaaaaaaaaaaaaaaaaaaa", kind=LogicalListKind.UNORDERED,
        member_fragment_ids=(fragment_ids[11], fragment_ids[12]),
    )
    child = LogicalListV3(
        list_id="list_bbbbbbbbbbbbbbbbbbbb", kind=LogicalListKind.ORDERED,
        member_fragment_ids=(fragment_ids[13], fragment_ids[14]),
        parent_list_id=parent.list_id, parent_item_fragment_id=fragment_ids[12], start_value=3,
    )
    features = {
        item: FlowSourceFeatures(source_order=index, logical_sequence_explicit=types[index] is SemanticType.FIGURE)
        for index, item in enumerate(fragment_ids)
    }
    image = next(item for item in extraction.raw_document.objects if isinstance(item, RawImage))
    metadata = BookMetadataV3(
        title_fragment_id=fragment_ids[0], author_fragment_ids=(fragment_ids[1],),
        language="vi", identifier="urn:bookforge:m5c", cover_reference=image.asset_reference,
    )
    return classifier, regions, (parent, child), features, metadata, image.asset_reference


def test_rich_docx_runs_actual_pipeline_twice_byte_identically(tmp_path: Path) -> None:
    source, image_path = tmp_path / "golden.docx", tmp_path / "pixel.png"
    _make_docx(source, image_path)
    classifier, regions, lists, features, metadata, asset_reference = _fixture_truth(source, tmp_path)
    results = []
    for name in ("clean-a", "clean-b"):
        result = PipelineRunner().run(PipelineInput(
            source_docx=source, workspace_root=tmp_path / name,
            output_epub=tmp_path / name / "book.epub", metadata=metadata,
            semantic_classifier=classifier, structural_regions=regions,
            logical_lists=lists, source_features=features,
        ))
        assert result.assembly_readiness.ready
        assert result.structural_validation.status.value == "pass"
        assert result.source_state_sha256_before == result.source_state_sha256_after
        results.append(result)
    first, second = results
    assert first.extraction.document_id == second.extraction.document_id
    assert first.resolved_flow.model_dump_json() == second.resolved_flow.model_dump_json()
    assert first.book.model_dump_json() == second.book.model_dump_json()
    assert first.artifact.sha256 == second.artifact.sha256
    assert (tmp_path / "clean-a/book.epub").read_bytes() == (tmp_path / "clean-b/book.epub").read_bytes()
    resumed = PipelineRunner().run(PipelineInput(
        source_docx=source, workspace_root=tmp_path / "clean-a",
        output_epub=tmp_path / "clean-a/book.epub", metadata=metadata,
        semantic_classifier=classifier, structural_regions=regions,
        logical_lists=lists, source_features=features,
    ))
    assert resumed.semantic_report.reused == resumed.semantic_report.total_work_units
    assert resumed.flow_report.reused == resumed.flow_report.total_work_units
    assert resumed.book == first.book and resumed.artifact.sha256 == first.artifact.sha256
    with zipfile.ZipFile(tmp_path / "clean-a/book.epub") as package:
        names = package.namelist()
        assert names[0] == "mimetype"
        assert {"META-INF/container.xml", "EPUB/package.opf", "EPUB/nav.xhtml", "EPUB/styles.css"}.issubset(names)
        xhtml = "".join(package.read(name).decode() for name in names if name.endswith(".xhtml"))
        assert "Sách Việt &amp; XML &lt;đúng&gt;" in xhtml
        assert xhtml.count("<figure") == 2 and xhtml.count("<figcaption") == 1
        assert xhtml.count("<table") == 1 and xhtml.count("<ul") == 1
        assert xhtml.count('<ol start="3">') == 1
        for phrase in ("Mục lớn 1", "Mục lớn 2", "Mục con 2.1", "Mục con 2.2"):
            assert xhtml.count(phrase) == 1
        packaged_image = next(name for name in names if name.startswith("EPUB/images/"))
        extracted = first.extraction.workspace / asset_reference
        assert package.read(packaged_image) == extracted.read_bytes()
        assert not any("page" in Path(name).stem.lower() for name in names if name.endswith(".xhtml"))


def test_missing_explicit_truth_remains_unknown_and_blocks(tmp_path: Path) -> None:
    source, image_path = tmp_path / "unknown.docx", tmp_path / "pixel.png"
    _make_docx(source, image_path)
    classifier, regions, lists, features, metadata, _ = _fixture_truth(source, tmp_path)
    omitted = dict(classifier.semantic_by_unit)
    omitted.pop(next(iter(omitted)))
    with pytest.raises(AssemblyNotReadyError):
        PipelineRunner().run(PipelineInput(
            source_docx=source, workspace_root=tmp_path / "run",
            output_epub=tmp_path / "run/book.epub", metadata=metadata,
            semantic_classifier=ExplicitFixtureClassifier(omitted), structural_regions=regions,
            logical_lists=lists, source_features=features,
        ))


@pytest.mark.parametrize("continuity", [ContinuityType.JOIN_DIRECT, ContinuityType.JOIN_WITH_NEWLINE])
def test_reviewed_join_truth_traverses_actual_m4_m5_pipeline(tmp_path: Path, continuity: ContinuityType) -> None:
    source, image_path = tmp_path / "reviewed.docx", tmp_path / "pixel.png"
    _make_docx(source, image_path)
    classifier, regions, lists, features, metadata, _ = _fixture_truth(source, tmp_path)
    base = PipelineRunner().run(PipelineInput(
        source_docx=source, workspace_root=tmp_path / "base",
        output_epub=tmp_path / "base/book.epub", metadata=metadata,
        semantic_classifier=classifier, structural_regions=regions,
        logical_lists=lists, source_features=features,
    ))
    target = next(
        boundary for boundary in base.resolved_flow.boundaries
        if boundary.preceding_fragment_id == "sem_f000010"
        and boundary.following_fragment_id == "sem_f000011"
    )
    left = base.book.content.nodes[target.preceding_fragment_id]
    right = base.book.content.nodes[target.following_fragment_id]
    assert isinstance(left, TextSemanticNode) and isinstance(right, TextSemanticNode)
    suffix = "d" if continuity is ContinuityType.JOIN_DIRECT else "e"
    replacement = target.model_copy(update={
        "audit": target.audit.model_copy(update={"decision_id": f"fld_{suffix * 20}"}),
        "continuity": continuity,
        "source_references": (*left.source_references, *right.source_references),
    })
    review = FlowDecisionReview(
        review_id=f"fdr_{suffix * 20}", original_decision_id=target.audit.decision_id,
        status=ReviewStatus.REVIEWED_OVERRIDDEN,
        accepted_decision_id=replacement.audit.decision_id,
        reviewer=ResolverIdentity(name="m5c-review", kind=ResolverKind.HUMAN_REVIEW, version="1"),
        review_fingerprint=suffix * 64, created_at=EPOCH,
    )
    original_json = target.model_dump_json()
    reviewed = PipelineRunner().run(PipelineInput(
        source_docx=source, workspace_root=tmp_path / continuity.value,
        output_epub=tmp_path / continuity.value / "book.epub", metadata=metadata,
        semantic_classifier=classifier, structural_regions=regions,
        logical_lists=lists, source_features=features,
        flow_reviews=(AcceptedFlowReviewInput(review=review, replacement_decision=replacement),),
    ))
    preserved = next(item for item in reviewed.resolved_flow.boundaries if item.audit.decision_id == target.audit.decision_id)
    assert preserved.model_dump_json() == original_json
    assert reviewed.resolved_flow.decision_reviews == (review,)
    assert reviewed.resolved_flow.replacement_decisions == (replacement,)
    assert target.audit.decision_id not in reviewed.resolved_flow.unresolved_decision_ids
    assert any(item.operation is continuity for item in reviewed.book.continuity)
    assert reviewed.artifact.sha256 != base.artifact.sha256


def test_one_thousand_text_targets_cross_actual_pipeline(tmp_path: Path) -> None:
    source = tmp_path / "large.docx"
    doc = Document()
    doc.add_paragraph("Large Book")
    doc.add_paragraph("Chapter")
    for index in range(1000):
        text = "compre-" if index == 0 else "hensive" if index == 1 else f"Paragraph {index:04d}"
        doc.add_paragraph(text)
    doc.save(source)
    extraction = DocxExtractor().extract(source, tmp_path / "discovery-large")
    units = generate_work_units(extraction.raw_document)
    assert len(units) == 1002
    kinds = [SemanticType.BOOK_TITLE, SemanticType.CHAPTER_HEADING, *([SemanticType.PARAGRAPH] * 1000)]
    classifier = ExplicitFixtureClassifier({unit.work_unit_id: kind for unit, kind in zip(units, kinds, strict=True)})
    ids = tuple(f"sem_f{index:06d}" for index in range(1, 1003))
    regions = StructuralRegionAssignment(by_fragment_id={
        item: StructuralRegion.FRONT if index == 0 else StructuralRegion.BODY
        for index, item in enumerate(ids)
    })
    features = {
        item: FlowSourceFeatures(
            source_order=index,
            source_boundary_before=index >= 3,
            continuation_group_id="large-body" if index >= 2 else None,
        )
        for index, item in enumerate(ids)
    }
    result = PipelineRunner().run(PipelineInput(
        source_docx=source, workspace_root=tmp_path / "large-run",
        output_epub=tmp_path / "large-run/book.epub",
        metadata=BookMetadataV3(title_fragment_id=ids[0], language="en", identifier="large"),
        semantic_classifier=classifier, structural_regions=regions, source_features=features,
    ))
    assert result.semantic_report.total_work_units == 1002
    assert result.flow_report.total_fragments == 1002
    assert result.assembly_readiness.ready
    assert result.structural_validation.status.value == "pass"
    assert any(
        item.operation is ContinuityType.JOIN_REMOVE_TRAILING_HYPHEN
        for item in result.book.continuity
    )
    assert any(item.operation is ContinuityType.JOIN_WITH_SPACE for item in result.book.continuity)
    with zipfile.ZipFile(tmp_path / "large-run/book.epub") as package:
        rendered = "".join(
            package.read(name).decode() for name in package.namelist() if name.endswith(".xhtml")
        )
    assert "comprehensive" in rendered and "compre-" not in rendered
